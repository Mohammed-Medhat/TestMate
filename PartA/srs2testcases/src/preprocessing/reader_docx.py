# src/preprocessing/reader_docx.py
import json
import os
import re
from pathlib import Path
from docx import Document

# Basic cleaning utilities
_whitespace_re = re.compile(r"\s+")
_bullet_re = re.compile(r"^\s*(?:[\u2022\-\*\u2023]|\d+[\.\)])\s+")
_requirement_id_re = re.compile(r"^(REQ[_\-]?\d+|R-\d+|[A-Z]{2,}-\d+)", re.I)

def clean_text(text: str) -> str:
    if text is None:
        return ""
    text = text.replace('\xa0', ' ')  # non-breaking space
    text = _whitespace_re.sub(" ", text)
    return text.strip()

def paragraph_type(paragraph):
    """Heuristics to determine paragraph type"""
    txt = paragraph.text.strip()
    if not txt:
        return "empty"
    # heading by style name if available
    try:
        style_name = paragraph.style.name.lower()
    except Exception:
        style_name = ""
    if "heading" in style_name or "title" in style_name:
        return "heading"
    # list item by leading bullet/number
    if _bullet_re.match(txt):
        return "list_item"
    # short title-case heuristics
    if len(txt.split()) <= 6 and txt == txt.title():
        return "heading"
    return "paragraph"

def parse_docx(path, out_dir="data/raw", doc_id=None, parse_tables=True, keep_raw=True):
    """
    Read a .docx file and produce a JSON with 'blocks' and metadata.
    Each block: block_id, type, text, raw_text (optional), style, table_id (if table_row), row_index (if table_row)x
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(path)

    doc = Document(str(path))
    output = {
        "file_path": str(path),
        "doc_id": doc_id or path.stem,
        "num_paragraphs": len(doc.paragraphs),
        "blocks": []
    }

    block_counter = 0

    # 1) Parse paragraphs and lists (order follows doc.paragraphs)
    for i, p in enumerate(doc.paragraphs):
        raw = p.text or ""
        cleaned = clean_text(raw)
        if not cleaned:
            # still add empty blocks optionally or skip
            continue

        btype = paragraph_type(p)
        try:
            style_name = p.style.name
        except Exception:
            style_name = None

        # remove bullets from start for normalized text     (but keep raw_text)
        normalized_text = _bullet_re.sub("", cleaned)

        block = {
            "block_id": f"b{block_counter}",
            "type": btype,
            "text": normalized_text,
            "raw_text": raw if keep_raw else None,
            "style": style_name,
            "source": str(path),
            "paragraph_index": i
        }
        output["blocks"].append(block)
        block_counter += 1

    # 2) Parse tables (document.tables returns tables in reading order)
    if parse_tables and hasattr(doc, "tables"):
        for tidx, table in enumerate(doc.tables):
            # get rows as list of lists of cell texts
            rows = []
            for ridx, row in enumerate(table.rows):
                cells = [clean_text(cell.text) for cell in row.cells]
                rows.append(cells)

            # detect header row heuristic: first row with short uppercase-like tokens
            header = None
            if rows:
                first = rows[0]
                # header if many non-empty cells and short cell strings
                header_candidate = all(len(c.split()) <= 6 and c == c.title() for c in first if c)
                if header_candidate:
                    header = first

            # convert rows to blocks (skip header row as separate "table_header")
            start_index = 1 if header else 0
            for ridx in range(start_index, len(rows)):
                row_cells = rows[ridx]
                # create a concatenated text for the row (use " | " to preserve separation)
                row_text = " | ".join([c for c in row_cells if c])
                if not row_text:
                    continue
                block = {
                    "block_id": f"t{tidx}_r{ridx}",
                    "type": "table_row",
                    "table_id": f"table_{tidx}",
                    "row_index": ridx,
                    "text": row_text,
                    "raw_row": row_cells if keep_raw else None,
                    "header": header if (ridx == 1 and header) else None,
                    "source": str(path)
                }
                output["blocks"].append(block)
                block_counter += 1

    # Optional: attempt to detect explicit requirement IDs inside blocks and set metadata
    for b in output["blocks"]:
        m = _requirement_id_re.match(b["text"])
        if m:
            b["detected_req_id"] = m.group(0)

    # stats
    output["num_blocks"] = len(output["blocks"])
    # save json
    os.makedirs(out_dir, exist_ok=True)
    out_path = Path(out_dir) / f"{output['doc_id']}.json"
    with open(out_path, "w", encoding="utf-8") as fh:
        json.dump(output, fh, ensure_ascii=False, indent=2)
    return str(out_path)

# CLI usage
if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Read a .docx and save normalized JSON blocks.")
    parser.add_argument("docx", help="path to .docx file")
    parser.add_argument("--out", help="output directory (default=data/raw)", default="data/raw")
    parser.add_argument("--no-tables", action="store_true", help="do not parse tables")
    parser.add_argument("--no-raw", action="store_true", help="do not keep raw_text/raw_row")
    args = parser.parse_args()
    outp = parse_docx(args.docx, out_dir=args.out, parse_tables=not args.no_tables, keep_raw=not args.no_raw)
    print("Saved JSON to:", outp)
