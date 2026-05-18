"""
Convert requests_SRS.md  →  requests_SRS.docx

Needed because Part A's SRS pipeline reads DOCX/PDF, not Markdown.

Usage:
    pip install python-docx
    python generate_srs_docx.py
"""
from pathlib import Path
import re

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:
    raise SystemExit("Install python-docx first:  pip install python-docx")

HERE = Path(__file__).resolve().parent
MD   = HERE / "requests_SRS.md"
OUT  = HERE / "requests_SRS.docx"


def md_to_docx(md_text: str, out_path: Path) -> None:
    doc = Document()

    # ── Base styles ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    in_code = False
    code_buf: list[str] = []

    for raw in md_text.splitlines():
        line = raw.rstrip()

        # ── Fenced code block ──
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_buf = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue

        # ── Headings ──
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(m.group(2).strip(), level=level)
            continue

        # ── Horizontal rule ──
        if re.match(r"^[-*_]{3,}\s*$", line):
            p = doc.add_paragraph()
            p.add_run("─" * 60)
            continue

        # ── Bullet list ──
        m = re.match(r"^[-*]\s+(.+)$", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, m.group(1))
            continue

        # ── Plain paragraph ──
        if line.strip():
            p = doc.add_paragraph()
            _add_inline(p, line)
        else:
            doc.add_paragraph()  # blank line → spacer

    doc.save(out_path)


def _add_inline(p, text: str) -> None:
    """Render **bold** and `code` runs inline within a paragraph."""
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)  # emerald
        else:
            p.add_run(part)


if __name__ == "__main__":
    if not MD.exists():
        raise SystemExit(f"Markdown SRS not found: {MD}")
    md_to_docx(MD.read_text(encoding="utf-8"), OUT)
    print(f"Wrote {OUT}")
    print(f"Size: {OUT.stat().st_size:,} bytes")
